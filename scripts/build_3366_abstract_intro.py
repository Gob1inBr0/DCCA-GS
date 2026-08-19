# -*- coding: utf-8 -*-
"""Replace Abstract + Introduction contributions in the 3366 DCCA-GS docx.

Only paragraphs inside the Abstract and Introduction sections are changed;
the rest of the document (method, experiments, conclusion, references) is
preserved exactly as the user edited it.
"""
from docx import Document

SRC = ("/Users/chen/Desktop/3366-DCCA-GS：Decoder-Reproducible "
       "Content-Adaptive Compression for Anchor-Based 3D Gaussian "
       "Splatting.docx")

NEW_ABSTRACT = (
    "3D Gaussian Splatting (3DGS) achieves high-quality real-time novel-view "
    "rendering with explicit Gaussian primitives, but the per-primitive "
    "storage of geometry, covariance, opacity, and appearance grows rapidly "
    "with scene scale, making storage and transmission a deployment "
    "bottleneck for large-scale UAV and multi-view capture. This contribution "
    "describes DCCA-GS (Decoder-Reproducible Content-Adaptive Compression "
    "for Anchor-Based 3D Gaussian Splatting), built on the HAC++ anchor-hash "
    "entropy-coding framework. DCCA-GS keeps the HAC++ main pipeline "
    "(structured anchors, hash-grid context, conditional entropy model, "
    "arithmetic coding, G-PCC geometry) and adds four evaluated components: "
    "(1) decoder-reproducible content-adaptive formula quantization (I2), in "
    "which every quantization step is recomputed on both ends from a "
    "4-dimensional descriptor of decoder-reconstructable structural "
    "quantities, so no per-anchor step data is transmitted; (2) a "
    "training-only render-sensitivity supervision (I6) that steers the "
    "complexity network without changing the bitstream; (3) mixed-precision "
    "post-training quantization of the decoder MLPs, with a per-group "
    "bit-width assignment derived from per-MLP sensitivity ablations, "
    "charged to the reported size; and (4) training-side anchor sparsity "
    "with ADMM projection (SPA), which imposes a hard anchor budget during "
    "training and re-adapts the surviving anchors under that budget, "
    "defining a low-rate operating point (74% size reduction at 0.84 dB "
    "PSNR cost on Deep-Blending playroom at 110k). Components (1) and (2) "
    "are jointly designed; their interaction is analyzed in the ablation "
    "study, where I6 is critical at low-rate SPA points (+0.42 dB). On the "
    "4-28 UAV scene, the 110k-iteration operating point with MLP "
    "quantization reaches 28.82 dB PSNR at 5.49 MiB decode-required size, "
    "versus 28.31 dB at 6.95 MiB for the paper-reported HAC++ point on the "
    "same scene; on Deep Blending, playroom reaches 30.73 dB at 4.22 MiB "
    "and drjohnson 30.04 dB at 6.85 MiB. We also report quantitative "
    "analyses of codec efficiency (actual/estimated bits close to 1.0), a "
    "KL-redundancy lower bound for the scaling field, bitstream "
    "composition, rate-distortion behavior over lambda and training steps, "
    "and documented negative results (context conditioning below 2.3%, "
    "residual coding, and sensitivity side information). All bitstreams "
    "pass bit-exact standalone decoding verification."
)

NEW_INTRO = [
    "The two training-side modules are labeled I2 and I6 following our "
    "internal development taxonomy (I1 hierarchical context and I5 VQ were "
    "explored and are not part of the recommended configuration); the "
    "numbering is historical, and only the retained modules are described "
    "here.",
    "Our contributions are as follows.",
    "1. Decoder-reproducible content-adaptive formula quantization (I2). "
    "Quantization steps are generated on the fly from a 4-dimensional "
    "descriptor of decoder-reconstructable structural quantities—local "
    "anchor density, predicted scaling anisotropy, predicted offset energy, "
    "and active mask ratio. The bitstream carries no per-anchor step array, "
    "and training uses a straight-through estimator so that training, "
    "encoding, and decoding share the same quantization path. I2 is most "
    "effective when combined with I6 (Section 3.4); isolated-I2 conclusions "
    "are avoided at low-rate SPA operating points.",
    "2. Render-sensitivity supervision (I6). The complexity network is "
    "trained to match a bounded, exponentially smoothed target derived from "
    "render-loss gradient norms; this affects training only and adds zero "
    "side information. Its high-rate gain is modest (about 0.1 dB), but it "
    "becomes critical in the low-rate SPA regime, where removing it reduces "
    "PSNR by 0.42 dB (see Section 3.4).",
    "3. Mixed-precision MLP weight quantization. We apply per-output-channel "
    "symmetric post-training quantization with static arithmetic coding. "
    "The bit-width assignment is not uniform: the complexity and deform "
    "networks are quantized to 8 bits, while the opacity, covariance, color, "
    "and grid networks remain at 16 bits. This assignment is derived from "
    "per-MLP sensitivity ablation: 8-bit opacity directly degrades "
    "rendering, and 8-bit grid degrades the entropy model and increases "
    "attribute rate, whereas complexity/deform tolerate 8 bits with "
    "negligible loss.",
    "4. Training-side anchor sparsity (SPA). Unlike encode-side top-k "
    "pruning, SPA imposes a hard anchor budget during training via ADMM "
    "projection and a linear budget ramp anchored to the historical maximum "
    "anchor count, allowing surviving anchors to re-adapt under the budget. "
    "At equal anchor count it outperforms encode-side top-k pruning by "
    "5.5 dB, and the resulting operating point trades 0.84 dB PSNR for a "
    "74% size reduction relative to the non-SPA model (see Section 3.4).",
    "In the implementation and experiments, components (1) and (2) are "
    "referred to as I2 and I6, respectively; they are jointly designed, and "
    "their interaction is analyzed in the ablation study. All experiments "
    "follow a strict bitstream contract with standalone decoding "
    "verification and quantitative analyses of where the bits go and how "
    "much redundancy remains; this constitutes the evaluation protocol of "
    "this work rather than a separate technical contribution.",
]


def set_paragraph_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def main():
    doc = Document(SRC)
    paras = doc.paragraphs
    # Abstract body: first non-empty paragraph after "Abstract"
    for i, p in enumerate(paras):
        if p.text.strip() == "Abstract":
            set_paragraph_text(paras[i + 1], NEW_ABSTRACT)
            break
    # Introduction contributions block
    target = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith("Our contributions are as follows."):
            target = i
            break
    if target is None:
        raise RuntimeError("Introduction contributions paragraph not found")
    old = paras[target]
    for text in NEW_INTRO:
        old.insert_paragraph_before(text)
    old._element.getparent().remove(old._element)
    doc.save(SRC)
    print("saved", SRC)


if __name__ == "__main__":
    main()
