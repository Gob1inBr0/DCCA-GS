# -*- coding: utf-8 -*-
"""Replace I2/I6 labels with FQ/RSS in the 3366 DCCA-GS proposal docx."""
from docx import Document

SRC = ("/Users/chen/Desktop/3366-DCCA-GS：Decoder-Reproducible "
       "Content-Adaptive Compression for Anchor-Based 3D Gaussian "
       "Splatting.docx")

REPLACEMENTS = [
    ("I2/I6", "FQ/RSS"),
    ("I2+I6", "FQ+RSS"),
    ("I2", "FQ"),
    ("I6", "RSS"),
]


def replace_in_runs(paragraph):
    changed = False
    for run in paragraph.runs:
        t = run.text
        for old, new in REPLACEMENTS:
            if old in t:
                t = t.replace(old, new)
                changed = True
        if t != run.text:
            run.text = t
    return changed


def main():
    doc = Document(SRC)
    # 1) remove the internal-taxonomy explanation paragraph
    for p in list(doc.paragraphs):
        if p.text.strip().startswith(
            "The two training-side modules are labeled I2 and I6"
        ):
            p._element.getparent().remove(p._element)
    # 2) replace labels in all body paragraphs
    for p in doc.paragraphs:
        replace_in_runs(p)
    # 3) replace labels in all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs(p)
    doc.save(SRC)
    print("saved", SRC)


if __name__ == "__main__":
    main()
