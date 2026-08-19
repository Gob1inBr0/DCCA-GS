# -*- coding: utf-8 -*-
"""Fix small accuracy issues in the 3366 DCCA-GS proposal docx."""
from docx import Document

SRC = ("/Users/chen/Desktop/3366-DCCA-GS：Decoder-Reproducible "
       "Content-Adaptive Compression for Anchor-Based 3D Gaussian "
       "Splatting.docx")

NEW_CONTRIB = (
    "Our contributions are as follows. (1) Decoder-reproducible formula "
    "quantization (FQ): quantization steps are generated on the fly from a "
    "4-dimensional descriptor of decoder-reconstructable structural "
    "quantities—local anchor density, predicted scaling anisotropy, "
    "predicted offset energy, and active mask ratio—so the bitstream "
    "carries no per-anchor step array, and training uses a straight-through "
    "estimator shared by training, encoding, and decoding. (2) "
    "Render-sensitivity supervision (RSS): the complexity network is "
    "trained to match a bounded, exponentially smoothed target derived from "
    "render-loss gradient norms; RSS affects training only and adds zero "
    "side information. Its high-rate gain is about 0.1 dB, but it becomes "
    "critical in the low-rate SPA regime, where removing it reduces PSNR by "
    "0.42 dB. (3) Mixed-precision MLP weight quantization: "
    "per-output-channel symmetric quantization plus static arithmetic "
    "coding, with complexity/deform at 8 bit and the remaining groups at 16 "
    "bit; the compressed payload is included in the reported size. (4) "
    "Training-side anchor sparsity (SPA): an ADMM hard budget that prunes "
    "anchors by the top-k of (soft score + multiplier), with a linear "
    "budget ramp; at equal anchor count it outperforms encode-side top-k "
    "pruning by 5.5 dB, and the resulting operating point trades 0.84 dB "
    "PSNR for a 74% size reduction. All experiments follow a strict "
    "bitstream contract with standalone decoding verification and "
    "quantitative analyses of where the bits go and how much redundancy "
    "remains; this constitutes the evaluation protocol of this work rather "
    "than a separate technical contribution."
)


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def main():
    doc = Document(SRC)
    for p in doc.paragraphs:
        t = p.text
        if t.startswith("3D Gaussian Splatting (3DGS) achieves"):
            set_text(p, t.replace("describes DCCA-GS , built",
                                  "describes DCCA-GS, built")
                      .replace("a  Deep-Blending study", "a Deep-Blending study"))
        elif t.startswith("Our contributions are as follows."):
            set_text(p, NEW_CONTRIB)
        elif "best average rank of 3.50" in t:
            set_text(p, t.replace(
                "the best average rank of 3.50",
                "the best average ranks among all methods (4.00 for the "
                "no-SPA point and 4.50 for the SPA point)"))
        elif "4-28 sence" in t:
            set_text(p, t.replace("4-28 sence", "4-28 scene"))
        elif "an 4-dimensional descriptor" in t:
            set_text(p, t.replace("an 4-dimensional descriptor",
                                  "a 4-dimensional descriptor"))
    for ti, table in enumerate(doc.tables):
        if ti == 1:
            mapping = {"ours": "ours (SPA)", "ours(o SPA)": "ours (w/o SPA)"}
        elif ti == 2:
            mapping = {"ours": "ours (w/o SPA)", "ours (o SPA)": "ours (SPA)"}
        else:
            mapping = {}
        for row in table.rows:
            for cell in row.cells:
                c = cell.text.strip()
                if c in mapping:
                    for p in cell.paragraphs:
                        set_text(p, mapping[c])
                elif c == "1" and ti == 5:
                    for p in cell.paragraphs:
                        set_text(p, "100%")
    doc.save(SRC)
    print("saved", SRC)


if __name__ == "__main__":
    main()
